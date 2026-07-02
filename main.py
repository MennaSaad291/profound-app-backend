import os
import bcrypt
import io
import json
import re
import uuid
import shutil
import pandas as pd
from difflib import SequenceMatcher
from typing import List, Optional
from datetime import datetime
from dotenv import load_dotenv

# Load .env FIRST before any service imports that read env vars
load_dotenv()

from fastapi import FastAPI, HTTPException, Depends, Form, File, UploadFile
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy import func
from sqlalchemy.orm import Session
from docx import Document
from groq import Groq
from typing import List

from database import engine, Base, get_db, SessionLocal
from models import (
    UserDB, CourseDB, StudentDB, PublicationDB, ProjectDB, InterestDB,
    GraduationProjectDB, LiteraturePaperDB,
    ExamDB, QuestionDB, SubmissionDB, PerformanceDB, ErrorAnalysisDB,
    AssignmentDB, GradeUpdate, FinalizeRequest, LectureSlotDB
)
from schemas import (
    UserCreate, UserLogin, UserUpdate, LectureRequest, CourseResponse,
    ExamRequest, ExamResponse, Question, AssignmentCreate, CourseCreate,
    ChangePasswordRequest, VerifyPasswordRequest
)
from routes import analysis, lecture
from routes import courses as courses_router
from file_utils import extract_text
from services.grading_service import perform_nlp_grading
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch

try:
    from services.pptx_service import create_pptx
except ImportError:
    create_pptx = None

load_dotenv()
Base.metadata.create_all(bind=engine)

app = FastAPI(title="Profound Academic API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

groq_client = Groq(api_key=os.getenv("GROQ_API_KEY_ANALYSIS"))

UPLOAD_DIR = "uploads/assignments"
os.makedirs(UPLOAD_DIR, exist_ok=True)
GRAD_DOCS_DIR = "uploads/grad_projects"
os.makedirs(GRAD_DOCS_DIR, exist_ok=True)


@app.on_event("startup")
def sync_course_student_counts():
    """Re-sync the cached CourseDB.students column with the live StudentDB count
    on every startup so stale values from past bugs are corrected immediately."""
    db = SessionLocal()
    try:
        courses = db.query(CourseDB).all()
        for course in courses:
            live_count = (
                db.query(func.count(StudentDB.id))
                .filter(StudentDB.course_id == course.id)
                .scalar() or 0
            )
            if course.students != live_count:
                course.students = live_count
        db.commit()
    except Exception as e:
        print(f"[startup] student count sync failed: {e}")
    finally:
        db.close()


def clean_markdown(text: str) -> str:
    return re.sub(r'\*\*(.*?)\*\*', r'\1', str(text)).replace('*', '').strip()


# ── Auth ──────────────────────────────────────────────────────────────────────

@app.post("/register")
def register_user(user: UserCreate, db: Session = Depends(get_db)):
    if db.query(UserDB).filter(UserDB.email == user.email.lower()).first():
        raise HTTPException(status_code=400, detail="Email already registered")
    salt = bcrypt.gensalt()
    db.add(UserDB(
        full_name=user.full_name,
        email=user.email.lower(),
        password_hash=bcrypt.hashpw(user.password.encode(), salt).decode()
    ))
    db.commit()
    return {"message": "Success"}


@app.post("/login")
def login_user(user: UserLogin, db: Session = Depends(get_db)):
    db_user = db.query(UserDB).filter(UserDB.email == user.email.lower()).first()
    if not db_user or not bcrypt.checkpw(user.password.encode(), db_user.password_hash.encode()):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    return {"user": {"id": db_user.id, "name": db_user.full_name}}


@app.get("/profile/{user_id}")
def get_profile(user_id: int, db: Session = Depends(get_db)):
    user = db.query(UserDB).filter(UserDB.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    pubs = db.query(PublicationDB).filter(PublicationDB.user_id == user_id).all()
    courses = db.query(CourseDB).filter(CourseDB.user_id == user_id).all()
    grad_projects = db.query(GraduationProjectDB).filter(GraduationProjectDB.user_id == user_id).all()
    interests = db.query(InterestDB).filter(InterestDB.user_id == user_id).all()
    return {
        "id": user.id, "full_name": user.full_name, "email": user.email,
        "bio": user.bio, "department": user.department,
        "metrics": {
            "citations": sum(p.citations for p in pubs),
            "students": db.query(func.count(StudentDB.id)).filter(StudentDB.course_id.in_([c.id for c in courses])).scalar() or 0,
            "papers": len(pubs),
            "projects": len(grad_projects),
        },
        "publications": [
            {"id": p.id, "title": p.title, "journal": p.journal,
             "year": p.year, "citations": p.citations}
            for p in pubs
        ],
        "projects": [
            {"id": g.id, "title": g.title, "team": g.team,
             "academic_year": g.academic_year, "department": g.department,
             "has_document": bool(g.document_path)}
            for g in grad_projects
        ],
        "interests": [i.name for i in interests],
    }


@app.get("/research/{user_id}")
def get_research(user_id: int, db: Session = Depends(get_db)):
    """Return all research data for a user: publications, projects, interests, literature."""
    pubs     = db.query(PublicationDB).filter(PublicationDB.user_id == user_id).all()
    projects = db.query(ProjectDB).filter(ProjectDB.user_id == user_id).all()
    interests = db.query(InterestDB).filter(InterestDB.user_id == user_id).all()
    papers   = db.query(LiteraturePaperDB).filter(LiteraturePaperDB.user_id == user_id).all()

    in_progress  = sum(1 for p in projects if (p.status or '').lower() in ('ongoing', 'in progress', 'drafting'))
    under_review = sum(1 for p in projects if (p.status or '').lower() in ('under-review', 'under review', 'submitted'))

    # Upcoming deadlines — projects that have a deadline set, sorted by date
    from datetime import date as dt_date
    today = dt_date.today()
    deadlines = []
    for p in projects:
        if p.deadline:
            try:
                d = dt_date.fromisoformat(p.deadline)
                days_left = (d - today).days
                deadlines.append({
                    "id": p.id,
                    "title": p.title,
                    "date": p.deadline,
                    "days_left": days_left,
                })
            except ValueError:
                pass
    deadlines.sort(key=lambda x: x["days_left"])

    return {
        "publications": [
            {"id": p.id, "title": p.title, "journal": p.journal,
             "year": p.year, "citations": p.citations}
            for p in pubs
        ],
        "projects": [
            {"id": p.id, "title": p.title, "team": p.team,
             "year": p.year, "status": p.status,
             "deadline": p.deadline, "progress": p.progress or 0}
            for p in projects
        ],
        "interests": [i.name for i in interests],
        "literature": [
            {"id": lp.id, "title": lp.title,
             "read_status": lp.read_status, "citation_format": lp.citation_format}
            for lp in papers
        ],
        "deadlines": deadlines,
        "stats": {
            "active_projects":    len(projects),
            "in_progress":        in_progress,
            "under_review":       under_review,
            "total_publications": len(pubs),
            "total_citations":    sum(p.citations or 0 for p in pubs),
            "to_read":            sum(1 for lp in papers if lp.read_status == "to-read"),
        }
    }


# ── Publications CRUD ─────────────────────────────────────────────────────────

@app.put("/publications/{pub_id}")
def update_publication(pub_id: int, data: dict, db: Session = Depends(get_db)):
    pub = db.query(PublicationDB).filter(PublicationDB.id == pub_id).first()
    if not pub:
        raise HTTPException(status_code=404, detail="Publication not found")
    pub.title     = data.get("title", pub.title)
    pub.journal   = data.get("journal", pub.journal)
    pub.year      = data.get("year", pub.year)
    pub.citations = data.get("citations", pub.citations)
    db.commit()
    return {"message": "Publication updated"}


@app.delete("/publications/{pub_id}")
def delete_publication(pub_id: int, db: Session = Depends(get_db)):
    pub = db.query(PublicationDB).filter(PublicationDB.id == pub_id).first()
    if not pub:
        raise HTTPException(status_code=404, detail="Publication not found")
    db.delete(pub)
    db.commit()
    return {"message": "Publication deleted"}


# ── Research Projects CRUD ────────────────────────────────────────────────────

@app.put("/projects/{project_id}")
def update_project(project_id: int, data: dict, db: Session = Depends(get_db)):
    proj = db.query(ProjectDB).filter(ProjectDB.id == project_id).first()
    if not proj:
        raise HTTPException(status_code=404, detail="Project not found")
    proj.title    = data.get("title", proj.title)
    proj.team     = data.get("team", proj.team)
    proj.year     = data.get("year", proj.year)
    proj.status   = data.get("status", proj.status)
    proj.deadline = data.get("deadline", proj.deadline)
    proj.progress = int(data.get("progress", proj.progress or 0))
    db.commit()
    return {"message": "Project updated"}


# ── Literature Papers CRUD ────────────────────────────────────────────────────

@app.post("/literature-papers")
def add_literature_paper(data: dict, db: Session = Depends(get_db)):
    paper = LiteraturePaperDB(
        user_id=data["user_id"],
        title=data["title"],
        read_status=data.get("read_status", "to-read"),
        citation_format=data.get("citation_format", "APA"),
    )
    db.add(paper)
    db.commit()
    db.refresh(paper)
    return {"message": "Paper added", "id": paper.id}


@app.put("/literature-papers/{paper_id}")
def update_literature_paper(paper_id: int, data: dict, db: Session = Depends(get_db)):
    paper = db.query(LiteraturePaperDB).filter(LiteraturePaperDB.id == paper_id).first()
    if not paper:
        raise HTTPException(status_code=404, detail="Paper not found")
    paper.title          = data.get("title", paper.title)
    paper.read_status    = data.get("read_status", paper.read_status)
    paper.citation_format = data.get("citation_format", paper.citation_format)
    db.commit()
    return {"message": "Paper updated"}


@app.delete("/literature-papers/{paper_id}")
def delete_literature_paper(paper_id: int, db: Session = Depends(get_db)):
    paper = db.query(LiteraturePaperDB).filter(LiteraturePaperDB.id == paper_id).first()
    if not paper:
        raise HTTPException(status_code=404, detail="Paper not found")
    db.delete(paper)
    db.commit()
    return {"message": "Paper deleted"}


@app.delete("/projects/{project_id}")
def delete_project(project_id: int, db: Session = Depends(get_db)):
    proj = db.query(ProjectDB).filter(ProjectDB.id == project_id).first()
    if not proj:
        raise HTTPException(status_code=404, detail="Project not found")
    db.delete(proj)
    db.commit()
    return {"message": "Project deleted"}


# ── Graduation Projects CRUD ──────────────────────────────────────────────────

@app.get("/graduation-projects/{user_id}")
def get_graduation_projects(user_id: int, db: Session = Depends(get_db)):
    items = db.query(GraduationProjectDB).filter(GraduationProjectDB.user_id == user_id).all()
    return [
        {"id": g.id, "title": g.title, "team": g.team,
         "academic_year": g.academic_year, "department": g.department,
         "has_document": bool(g.document_path)}
        for g in items
    ]


@app.post("/graduation-projects")
async def create_graduation_project(
    user_id: int = Form(...),
    title: str = Form(...),
    team: str = Form(...),
    academic_year: str = Form(...),
    department: str = Form(...),
    document: Optional[UploadFile] = File(None),
    db: Session = Depends(get_db),
):
    doc_path = None
    if document and document.filename:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe = f"{user_id}_{ts}_{document.filename}".replace(" ", "_")
        doc_path = os.path.join(GRAD_DOCS_DIR, safe)
        with open(doc_path, "wb") as f:
            shutil.copyfileobj(document.file, f)

    proj = GraduationProjectDB(
        user_id=user_id, title=title, team=team,
        academic_year=academic_year, department=department,
        document_path=doc_path,
    )
    db.add(proj)
    db.commit()
    db.refresh(proj)
    return {"message": "Graduation project created", "id": proj.id}


@app.put("/graduation-projects/{project_id}")
async def update_graduation_project(
    project_id: int,
    title: str = Form(...),
    team: str = Form(...),
    academic_year: str = Form(...),
    department: str = Form(...),
    document: Optional[UploadFile] = File(None),
    db: Session = Depends(get_db),
):
    proj = db.query(GraduationProjectDB).filter(GraduationProjectDB.id == project_id).first()
    if not proj:
        raise HTTPException(status_code=404, detail="Graduation project not found")
    proj.title = title
    proj.team = team
    proj.academic_year = academic_year
    proj.department = department
    if document and document.filename:
        if proj.document_path and os.path.exists(proj.document_path):
            os.remove(proj.document_path)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe = f"{project_id}_{ts}_{document.filename}".replace(" ", "_")
        proj.document_path = os.path.join(GRAD_DOCS_DIR, safe)
        with open(proj.document_path, "wb") as f:
            shutil.copyfileobj(document.file, f)
    db.commit()
    return {"message": "Graduation project updated"}


@app.delete("/graduation-projects/{project_id}")
def delete_graduation_project(project_id: int, db: Session = Depends(get_db)):
    proj = db.query(GraduationProjectDB).filter(GraduationProjectDB.id == project_id).first()
    if not proj:
        raise HTTPException(status_code=404, detail="Graduation project not found")
    if proj.document_path and os.path.exists(proj.document_path):
        os.remove(proj.document_path)
    db.delete(proj)
    db.commit()
    return {"message": "Graduation project deleted"}


@app.get("/graduation-projects/{project_id}/view")
def view_graduation_document(project_id: int, db: Session = Depends(get_db)):
    """Serves the PDF inline so the browser opens it in a new tab."""
    proj = db.query(GraduationProjectDB).filter(GraduationProjectDB.id == project_id).first()
    if not proj or not proj.document_path or not os.path.exists(proj.document_path):
        raise HTTPException(status_code=404, detail="Document not found")
    filename = os.path.basename(proj.document_path)
    return StreamingResponse(
        open(proj.document_path, "rb"),
        media_type="application/pdf",
        headers={"Content-Disposition": f"inline; filename={filename}"},
    )


@app.get("/graduation-projects/{project_id}/download")
def download_graduation_document(project_id: int, db: Session = Depends(get_db)):
    """Forces a file download."""
    proj = db.query(GraduationProjectDB).filter(GraduationProjectDB.id == project_id).first()
    if not proj or not proj.document_path or not os.path.exists(proj.document_path):
        raise HTTPException(status_code=404, detail="Document not found")
    filename = os.path.basename(proj.document_path)
    return StreamingResponse(
        open(proj.document_path, "rb"),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@app.put("/profile/{user_id}")
def update_profile(user_id: int, data: UserUpdate, db: Session = Depends(get_db)):
    user = db.query(UserDB).filter(UserDB.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    user.full_name = data.full_name
    user.bio = data.bio
    user.department = data.department
    db.commit()
    return {"message": "Profile updated successfully"}


@app.post("/verify-password")
def verify_password(data: VerifyPasswordRequest, db: Session = Depends(get_db)):
    user = db.query(UserDB).filter(UserDB.id == data.user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    if not bcrypt.checkpw(data.password.encode(), user.password_hash.encode()):
        raise HTTPException(status_code=401, detail="Incorrect password")
    return {"message": "Password verified"}


@app.post("/change-password")
def change_password(data: ChangePasswordRequest, db: Session = Depends(get_db)):
    user = db.query(UserDB).filter(UserDB.id == data.user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    if not bcrypt.checkpw(data.current_password.encode(), user.password_hash.encode()):
        raise HTTPException(status_code=401, detail="Current password is incorrect")
    if data.current_password == data.new_password:
        raise HTTPException(status_code=400, detail="New password must differ from current")
    if len(data.new_password) < 6:
        raise HTTPException(status_code=400, detail="Password must be at least 6 characters")
    salt = bcrypt.gensalt()
    user.password_hash = bcrypt.hashpw(data.new_password.encode(), salt).decode()
    db.commit()
    return {"message": "Password updated successfully"}


# ── Dashboard ─────────────────────────────────────────────────────────────────

@app.get("/professors/{user_id}/courses", response_model=List[CourseResponse])
def get_courses(user_id: int, db: Session = Depends(get_db)):
    return db.query(CourseDB).filter(CourseDB.user_id == user_id).all()


@app.get("/dashboard-stats/{user_id}")
def get_dashboard_stats(user_id: int, db: Session = Depends(get_db)):
    # Count all courses (all statuses — professors want to see their full picture)
    courses = db.query(CourseDB).filter(CourseDB.user_id == user_id).all()
    course_ids = [c.id for c in courses]
    active_course_ids = [c.id for c in courses if (c.status or '').lower() == 'active']

    if not course_ids:
        return {
            "class_average": 0.0, "average_trend": 0.0,
            "at_risk_count": 0, "pending_grading": 0,
            "total_students": 0, "total_courses": 0,
            "active_courses": 0,
        }

    # ── Total students: union of students table + performance table ──────────
    # Students can exist in StudentDB (enrolled via roster upload) OR only in
    # PerformanceDB (grades uploaded directly). Count the union so neither
    # source is missed.
    students_enrolled = (
        db.query(func.count(StudentDB.id))
        .filter(StudentDB.course_id.in_(course_ids))
        .scalar() or 0
    )
    # PerformanceDB.student_id is a FK to StudentDB, so students there are
    # always in StudentDB. But count distinct student_ids per course to catch
    # any orphaned performance records just in case.
    perf_student_count = (
        db.query(func.count(func.distinct(PerformanceDB.student_id)))
        .filter(PerformanceDB.course_id.in_(course_ids))
        .scalar() or 0
    )
    total_students = max(students_enrolled, perf_student_count)

    # ── Class average & at-risk: prefer PerformanceDB (actual grades),
    #    fall back to graded SubmissionDB entries when no performance data exists ──

    perf_records = (
        db.query(PerformanceDB)
        .filter(PerformanceDB.course_id.in_(course_ids))
        .all()
    )

    if perf_records:
        # Build per-student average from PerformanceDB
        student_perf: dict = {}
        for p in perf_records:
            student_perf.setdefault(p.student_id, []).append(p.grade)

        all_grades = [g for grades in student_perf.values() for g in grades]
        class_average = round(sum(all_grades) / len(all_grades), 1) if all_grades else 0.0

        at_risk_count = sum(
            1 for grades in student_perf.values()
            if (sum(grades) / len(grades)) < 70
        )
    else:
        # Fall back to graded submissions
        graded_subs = (
            db.query(SubmissionDB)
            .join(AssignmentDB, SubmissionDB.assignment_id == AssignmentDB.id)
            .filter(
                AssignmentDB.course_id.in_(course_ids),
                SubmissionDB.status == "graded",
                SubmissionDB.ai_grade.isnot(None),
            ).all()
        )
        grades = [s.ai_grade for s in graded_subs]
        class_average = round(sum(grades) / len(grades), 1) if grades else 0.0

        student_grade_map: dict = {}
        for s in graded_subs:
            key = s.student_name or f"sub_{s.id}"
            student_grade_map.setdefault(key, []).append(s.ai_grade)
        at_risk_count = sum(
            1 for g in student_grade_map.values() if sum(g) / len(g) < 70
        )

    pending_grading = (
        db.query(SubmissionDB)
        .join(AssignmentDB, SubmissionDB.assignment_id == AssignmentDB.id)
        .filter(
            AssignmentDB.course_id.in_(course_ids),
            SubmissionDB.status.in_(["pending", "ready"]),
        ).count()
    )

    return {
        "class_average":   class_average,
        "average_trend":   0.0,
        "at_risk_count":   at_risk_count,
        "pending_grading": pending_grading,
        "total_students":  total_students,
        "total_courses":   len(courses),
        "active_courses":  len(active_course_ids),
    }


# ── Bloom's taxonomy definitions used in every prompt ────────────────────────
BLOOMS_DEFINITIONS = {
    "Remember":   "RECALL facts, terms, definitions, dates. Questions must ask students to RECITE or IDENTIFY information directly stated in the material. Use verbs: define, list, name, recall, state, identify.",
    "Understand": "EXPLAIN concepts in their own words, SUMMARIZE ideas, INTERPRET meaning, CLASSIFY examples. Questions must ask students to DESCRIBE, EXPLAIN, or PARAPHRASE — not just recall. Use verbs: explain, summarise, classify, describe, interpret, compare.",
    "Apply":      "USE knowledge to SOLVE problems in NEW situations. Questions must present a SCENARIO or problem the student must solve by applying a rule, formula, or concept. Use verbs: solve, demonstrate, use, calculate, apply, show how.",
    "Analyze":    "BREAK DOWN information into parts, find PATTERNS, CAUSES, or RELATIONSHIPS. Questions must ask students to EXAMINE, COMPARE, CONTRAST, or DISTINGUISH components. Use verbs: analyse, compare, contrast, differentiate, examine, break down.",
    "Evaluate":   "JUDGE the value of information using CRITERIA, DEFEND a position, CRITIQUE a solution. Questions must ask students to ASSESS, ARGUE, JUSTIFY, or CRITIQUE with reasons. Use verbs: evaluate, justify, assess, critique, defend, recommend, judge.",
    "Create":     "DESIGN or PRODUCE something NEW by combining ideas. Questions must ask students to CONSTRUCT, DESIGN, PROPOSE, or PLAN something original. Use verbs: design, create, construct, propose, develop, formulate, plan.",
}


@app.post("/exams/generate", response_model=ExamResponse)
async def generate_exam(request: ExamRequest, db: Session = Depends(get_db)):
    blooms_def = BLOOMS_DEFINITIONS.get(request.blooms_level, "")

    prompt = f"""You are a university professor generating EXACTLY {request.number_of_questions} {request.question_type} exam questions.
Topic: {request.topic}
Difficulty: {request.difficulty}

BLOOM'S TAXONOMY LEVEL — {request.blooms_level.upper()}:
{blooms_def}

CRITICAL: Every question MUST operate at the {request.blooms_level} cognitive level as described above.
Questions at the WRONG level will be rejected. Do not generate lower-level questions.

QUESTION TYPE RULES:
- question_type for every question MUST be "{request.question_type}"
- MCQ: options must be a list of exactly 4 strings. correct_answer must be the FULL TEXT of the correct option.
- Essay: options must be null. correct_answer must be a detailed model answer string.

Return ONLY JSON with key "questions" containing exactly {request.number_of_questions} objects:
"question_text", "question_type", "options", "correct_answer", "explanation", "difficulty"

The explanation field must describe WHY this question tests {request.blooms_level}-level thinking."""

    try:
        completion = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": "You are a JSON-only academic exam generator. Strictly follow Bloom's taxonomy level instructions."},
                {"role": "user", "content": prompt}
            ],
            response_format={"type": "json_object"}
        )
        ai_data = json.loads(completion.choices[0].message.content)
        all_questions = ai_data.get("questions", [])[:request.number_of_questions]

        # Reuse an existing exam if provided (so all batches share one exam_id)
        existing_id = getattr(request, "existing_exam_id", None)
        if existing_id and db.query(ExamDB).filter(ExamDB.id == existing_id).first():
            exam_id = existing_id
        else:
            exam_id = str(uuid.uuid4())[:8]
            db.add(ExamDB(id=exam_id, course_id=request.course_id, title=f"Assessment: {request.topic}"))

        generated_questions = []
        for q in all_questions:
            q_options = q.get("options")
            raw_answer = q.get("correct_answer") or q.get("answer")
            if isinstance(raw_answer, int) and q_options and 0 <= raw_answer < len(q_options):
                q_answer = str(q_options[raw_answer])
            elif isinstance(raw_answer, str) and len(raw_answer) == 1 and raw_answer.upper() in "ABCD" and q_options:
                idx = ord(raw_answer.upper()) - ord('A')
                q_answer = str(q_options[idx]) if 0 <= idx < len(q_options) else raw_answer
            else:
                q_answer = str(raw_answer) if raw_answer is not None else ""

            q_text = q.get("question_text") or q.get("question")
            q_type = q.get("question_type") or request.question_type
            q_difficulty = q.get("difficulty") or request.difficulty

            db.add(QuestionDB(
                id=str(uuid.uuid4())[:8], exam_id=exam_id, question_text=q_text,
                question_type=q_type, options=q_options, blooms_level=request.blooms_level,
                difficulty=q_difficulty, correct_answer=q_answer,
                explanation=q.get("explanation", "")
            ))
            generated_questions.append(Question(
                question_text=q_text, options=q_options, correct_answer=q_answer,
                explanation=q.get("explanation", ""), difficulty=q_difficulty, question_type=q_type
            ))

        db.commit()
        return {"exam_id": exam_id, "questions": generated_questions}

    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/exams/export-word/{exam_id}")
async def export_exam_word(exam_id: str, db: Session = Depends(get_db)):
    exam = db.query(ExamDB).filter(ExamDB.id == exam_id).first()
    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")
    questions = db.query(QuestionDB).filter(QuestionDB.exam_id == exam_id).all()

    doc = Document()
    doc.add_heading(exam.title, 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph(f"Total Questions: {len(questions)}")
    doc.add_paragraph("")

    for i, q in enumerate(questions, 1):
        p = doc.add_paragraph()
        p.add_run(f"Q{i} [{q.question_type} | {q.difficulty}]: {q.question_text}").bold = True
        if q.question_type == "MCQ" and q.options:
            for opt in q.options:
                doc.add_paragraph(f"   [ ] {opt}")
        doc.add_paragraph("_" * 50)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=Exam_{exam_id}.docx"}
    )


@app.post("/exams/{exam_id}/variations")
async def generate_exam_variations(
    exam_id: str,
    data: dict,
    db: Session = Depends(get_db)
):
    """
    Generate N shuffled variations to minimise cheating.
    Each version has a different question order AND different MCQ option order.
    Each version is seeded differently so they are truly distinct.
    """
    import random, zipfile
    from copy import deepcopy

    num_variations = max(2, min(int(data.get("num_variations", 3)), 10))

    exam = db.query(ExamDB).filter(ExamDB.id == exam_id).first()
    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")

    questions = db.query(QuestionDB).filter(QuestionDB.exam_id == exam_id).all()
    if not questions:
        raise HTTPException(status_code=404, detail="No questions found for this exam")

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for v in range(1, num_variations + 1):
            # Different seed per version guarantees distinct shuffles
            rng = random.Random(v * 997 + len(questions) * 31)
            qs = deepcopy(questions)
            rng.shuffle(qs)

            doc = Document()
            doc.add_heading(f"{exam.title} — Version {v}", 0)
            doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}  |  Version {v} of {num_variations}")
            doc.add_paragraph(f"Total Questions: {len(qs)}")
            doc.add_paragraph("Student Name: ____________________________   ID: ____________")
            doc.add_paragraph("")

            answer_key = []
            labels = ["A", "B", "C", "D"]

            for i, q in enumerate(qs, 1):
                p = doc.add_paragraph()
                p.add_run(f"Q{i}. [{q.difficulty} | {getattr(q, 'blooms_level', '')}]  {q.question_text}").bold = True

                if q.question_type == "MCQ" and q.options:
                    opts = list(q.options)
                    correct_text = q.correct_answer
                    rng.shuffle(opts)   # unique shuffle per version
                    correct_label = "?"
                    for j, opt in enumerate(opts):
                        lbl = labels[j] if j < len(labels) else str(j + 1)
                        doc.add_paragraph(f"   {lbl}. {opt}")
                        if opt == correct_text:
                            correct_label = lbl
                    answer_key.append(f"Q{i}: {correct_label}")
                else:
                    doc.add_paragraph("Answer: _______________________________________________")
                    answer_key.append(f"Q{i}: {(q.correct_answer or '')[:100]}")

                doc.add_paragraph("")

            doc.add_page_break()
            doc.add_heading(f"ANSWER KEY — Version {v}  (Professor Copy — Do Not Distribute)", level=1)
            for ak in answer_key:
                doc.add_paragraph(ak)

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            zf.writestr(f"Exam_Version_{v:02d}.docx", buf.read())

    zip_buffer.seek(0)
    topic_clean = exam.title.replace(" ", "_")[:40]
    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename={topic_clean}_Variations.zip"}
    )


@app.post("/exams/generate-from-content")
async def generate_exam_from_content(
    topic: str = Form(...),
    number_of_questions: int = Form(5),
    difficulty: str = Form("Medium"),
    blooms_level: str = Form("Apply"),
    question_type: str = Form("MCQ"),
    mcq_percentage: int = Form(70),
    existing_exam_id: str = Form(None),     # if set, append questions to this exam
    content_file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    """Generate exam from uploaded content. Supports Mix, all Bloom's levels, and multi-batch."""
    content = await content_file.read()
    ext = content_file.filename.lower().split(".")[-1]
    temp_path = f"temp_exam_{uuid.uuid4()}.{ext}"
    with open(temp_path, "wb") as f:
        f.write(content)
    try:
        extracted = extract_text(temp_path)
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)

    content_excerpt = extracted[:3000] if len(extracted) > 3000 else extracted
    blooms_def = BLOOMS_DEFINITIONS.get(blooms_level, "")

    # Reuse existing exam or create new one
    if existing_exam_id:
        exam_id = existing_exam_id
        if not db.query(ExamDB).filter(ExamDB.id == exam_id).first():
            db.add(ExamDB(id=exam_id, title=f"Assessment: {topic}"))
    else:
        exam_id = str(uuid.uuid4())[:8]
        db.add(ExamDB(id=exam_id, title=f"Assessment: {topic}"))

    all_questions = []

    # Resolve Mix distribution
    if question_type == "Mix":
        mcq_count   = max(1, round(number_of_questions * mcq_percentage / 100))
        essay_count = max(1, number_of_questions - mcq_count)
        type_batches = [("MCQ", mcq_count), ("Essay", essay_count)]
    else:
        type_batches = [(question_type, number_of_questions)]

    for q_type, q_count in type_batches:
        if q_count == 0:
            continue
        prompt = f"""You are a university professor generating EXACTLY {q_count} {q_type} exam questions
based on the following lecture material about "{topic}".
Difficulty: {difficulty}

BLOOM'S TAXONOMY LEVEL — {blooms_level.upper()}:
{blooms_def}

CRITICAL: Every question MUST operate at the {blooms_level} cognitive level. Do not generate lower-level questions.

Lecture Material:
{content_excerpt}

RULES:
- Generate EXACTLY {q_count} questions derived from the provided material.
- question_type for every question MUST be "{q_type}"
- MCQ: options must be a list of exactly 4 strings. correct_answer must be the FULL TEXT of the correct option.
- Essay: options must be null. correct_answer must be a detailed model answer.

Return ONLY JSON with key "questions":
"question_text", "question_type", "options", "correct_answer", "explanation", "difficulty"
"""
        completion = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": "You are a JSON-only academic exam generator. Strictly follow Bloom's taxonomy instructions."},
                {"role": "user", "content": prompt}
            ],
            response_format={"type": "json_object"}
        )
        batch = json.loads(completion.choices[0].message.content).get("questions", [])[:q_count]

        for q in batch:
            q_options = q.get("options")
            raw_answer = q.get("correct_answer") or q.get("answer")
            if isinstance(raw_answer, int) and q_options and 0 <= raw_answer < len(q_options):
                q_answer = str(q_options[raw_answer])
            elif isinstance(raw_answer, str) and len(raw_answer) == 1 and raw_answer.upper() in "ABCD" and q_options:
                idx = ord(raw_answer.upper()) - ord('A')
                q_answer = str(q_options[idx]) if 0 <= idx < len(q_options) else raw_answer
            else:
                q_answer = str(raw_answer) if raw_answer is not None else ""

            q_text = q.get("question_text") or q.get("question")
            q_diff = q.get("difficulty") or difficulty

            db.add(QuestionDB(
                id=str(uuid.uuid4())[:8], exam_id=exam_id, question_text=q_text,
                question_type=q_type, options=q_options, blooms_level=blooms_level,
                difficulty=q_diff, correct_answer=q_answer,
                explanation=q.get("explanation", "")
            ))
            all_questions.append(Question(
                question_text=q_text, options=q_options, correct_answer=q_answer,
                explanation=q.get("explanation", ""), difficulty=q_diff, question_type=q_type
            ))

    try:
        db.commit()
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))

    return {"exam_id": exam_id, "questions": all_questions}


# ── Lecture / PPTX ───────────────────────────────────────────────────────────

def _build_pptx_response(data: dict):
    if create_pptx is None:
        raise HTTPException(status_code=500, detail="Install python-pptx to enable this feature.")
    if not data.get("slides"):
        raise HTTPException(status_code=400, detail="No slides provided")
    try:
        return StreamingResponse(
            create_pptx(data),
            media_type="application/vnd.openxmlformats-officedocument.presentationml.presentation",
            headers={"Content-Disposition": "attachment; filename=lecture.pptx"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/export-pptx")
async def export_pptx(data: dict):
    return _build_pptx_response(data)


@app.post("/export-pptx")
async def export_pptx_alias(data: dict):
    return _build_pptx_response(data)


# ── Assignments ───────────────────────────────────────────────────────────────
# ── Text Processing Utilities for Plagiarism ────────────────────────────

def clean_text(text: str) -> str:
    """Clean text for comparison"""
    if not text:
        return ""
    text = text.lower()
    text = re.sub(r'[^a-zA-Z0-9\s]', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def calculate_similarity(text1: str, text2: str) -> float:
    """Calculate similarity between two texts"""
    t1 = clean_text(text1)
    t2 = clean_text(text2)

    if not t1 or not t2:
        return 0.0

    return SequenceMatcher(None, t1, t2).ratio() * 100


async def check_plagiarism_ai_direct(student_text: str, user_id: int, db: Session):
    """
    Check plagiarism against PRIVATE submissions belonging to the specific user only
    """
    if not student_text:
        return {
            "plagiarism_score": 0,
            "matches": [],
            "message": "No student text provided"
        }

    existing_submissions = db.query(SubmissionDB).filter(SubmissionDB.user_id == user_id).all()
    if not existing_submissions:
        return {
            "plagiarism_score": 0,
            "matches": [],
            "message": "No existing submissions to check against"
        }

    matches = []
    highest_match = 0
    total_similarity = 0
    count = 0

    for sub in existing_submissions:
        if sub.essay_content and sub.id:
            similarity = calculate_similarity(student_text, sub.essay_content)

            if similarity > 15:
                matches.append({
                    "student_name": sub.student_name,
                    "similarity": similarity,
                    "submission_id": sub.id,
                    "content_preview": sub.essay_content[:200] + "..." if len(
                        sub.essay_content) > 200 else sub.essay_content,
                    "grade": sub.ai_grade
                })

                if similarity > highest_match:
                    highest_match = similarity

                total_similarity += similarity
                count += 1

    avg_similarity = total_similarity / count if count > 0 else 0

    if highest_match > 0:
        plagiarism_score = int((highest_match * 0.6) + (avg_similarity * 0.4))
        plagiarism_score = min(100, plagiarism_score)
    else:
        plagiarism_score = 0

    print(
        f"🔍 Plagiarism check (isolated): Found {len(matches)} user-specific matches, highest: {highest_match}%, avg: {avg_similarity}%, score: {plagiarism_score}%")
    return {
        "plagiarism_score": plagiarism_score,
        "highest_match": highest_match,
        "average_match": avg_similarity,
        "matches": sorted(matches, key=lambda x: x['similarity'], reverse=True),
        "total_checked": len(existing_submissions),
        "total_matches": len(matches),
        "message": f"Found {len(matches)} similar submissions out of {len(existing_submissions)} checked."
    }


# ── Submission & Assignment Routing ─────────────────────────────────────

@app.get("/submissions")
async def get_submissions(assignment_id: int, db: Session = Depends(get_db)):
    """Fetches submissions sorted by the latest record first."""
    return db.query(SubmissionDB).filter(
        SubmissionDB.assignment_id == assignment_id
    ).order_by(SubmissionDB.id.desc()).all()


@app.post("/assignments")
async def create_assignment(
        assignment_name: str = Form(...),
        course_id: int = Form(...),
        assignment_question: Optional[str] = Form(None),
        assignment_file: Optional[UploadFile] = File(None),
        is_model_answer: bool = Form(...),
        model_answer: Optional[str] = Form(None),
        rubric: Optional[str] = Form(None),
        db: Session = Depends(get_db)
):
    file_local_path = None

    if assignment_file:
        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_name = f"{course_id}_{timestamp}_{assignment_file.filename}".replace(" ", "_")
            file_local_path = os.path.join(UPLOAD_DIR, safe_name)

            with open(file_local_path, "wb") as buffer:
                shutil.copyfileobj(assignment_file.file, buffer)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"File save error: {str(e)}")

    course = db.query(CourseDB).filter(CourseDB.id == course_id).first()
    if not course:
        raise HTTPException(status_code=400, detail="Course not found")

    new_assignment = AssignmentDB(
        assignment_name=assignment_name,
        course_id=course_id,
        assignment_question=assignment_question,
        assignment_file_path=file_local_path,
        model_answer=model_answer,
        rubric=rubric,
        is_model_answer=is_model_answer
    )

    db.add(new_assignment)
    db.commit()
    db.refresh(new_assignment)
    return {"message": "Success", "id": new_assignment.id}


@app.get("/assignments-with-submissions/{course_id}")
async def get_nested_assignments(course_id: int, db: Session = Depends(get_db)):
    assignments = db.query(AssignmentDB).filter(AssignmentDB.course_id == course_id).all()
    result = []
    for assign in assignments:
        subs = db.query(SubmissionDB).filter(SubmissionDB.assignment_id == assign.id).all()
        result.append({
            "id": assign.id,
            "assignment_name": assign.assignment_name,
            "model_answer": assign.model_answer,
            "rubric": assign.rubric,
            "is_model_answer": assign.is_model_answer,
            "submissions": [
                {
                    "id": s.id,
                    "student_name": s.student_name,
                    "ai_grade": s.ai_grade,
                    "plagiarism_score": s.plagiarism_score,
                    "status": s.status,
                    "grade_report": s.grade_report,
                    "essay_content": s.essay_content,
                    "submitted_at": s.submission_time.isoformat() if s.submission_time else None
                } for s in subs
            ]
        })
    return result


@app.post("/grade-submission/{assignment_id}")
async def grade_student(
        assignment_id: int,
        file: UploadFile = File(...),
        feedback_tone: Optional[str] = Form("formal"),
        db: Session = Depends(get_db)
):
    assign = db.query(AssignmentDB).filter(AssignmentDB.id == assignment_id).first()
    if not assign:
        raise HTTPException(status_code=404, detail="Assignment not found")

    if feedback_tone not in {"formal", "encouraging", "strict"}:
        feedback_tone = "formal"

    content = await file.read()
    ext = file.filename.lower().split(".")[-1]
    temp_path = f"temp_{uuid.uuid4()}.{ext}"
    with open(temp_path, "wb") as f:
        f.write(content)
    student_text = ""
    try:
        if ext == "docx":
            from docx import Document
            student_text = "\n".join(p.text for p in Document(temp_path).paragraphs)
        elif ext == "pdf":
            student_text = extract_text(temp_path)
        else:
            raise HTTPException(status_code=400, detail="Only PDF or DOCX files are allowed")
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)

    if assign.is_model_answer:
        mode = "MODEL"
        reference = f"### MODE: MODEL_ANSWER\nQUESTION:\n{assign.assignment_question}\n\nIDEAL ANSWER:\n{assign.model_answer}"
    else:
        mode = "RUBRIC"
        reference = f"### MODE: RUBRIC\nQUESTION:\n{assign.assignment_question}\n\nRUBRIC:\n{assign.rubric}"

    # Evaluate using NLP module
    ai_report = perform_nlp_grading(
        student_text=student_text,
        mode=mode,
        reference=reference,
        feedback_tone=feedback_tone,
    )

    course = db.query(CourseDB).filter(CourseDB.id == assign.course_id).first()
    owner_id = course.user_id if course else 1

    plagiarism_result = await check_plagiarism_ai_direct(
        student_text=student_text,
        user_id=owner_id,
        db=db
    )
    plagiarism_score = plagiarism_result.get("plagiarism_score", 0)
    plagiarism_matches = plagiarism_result.get("matches", [])

    ai_report["plagiarism"] = plagiarism_score
    ai_report["plagiarism_matches"] = plagiarism_matches

    new_sub = SubmissionDB(
        assignment_id=assignment_id,
        student_name=file.filename.split('.')[0],
        status="ready",
        ai_grade=ai_report.get('score_out_of_100', 0),
        plagiarism_score=plagiarism_score,
        grade_report=ai_report,
        essay_content=student_text,
    )
    db.add(new_sub)
    db.flush()
    db.refresh(new_sub)

    # Track structural errors for reporting profiles
    category_label_map = {
        "conceptual": "Conceptual", "structural": "Structural",
        "language": "Language", "completeness": "Completeness",
    }
    student_record = db.query(StudentDB).filter(
        StudentDB.name == new_sub.student_name,
        StudentDB.course_id == assign.course_id
    ).first()

    for key, description in (ai_report.get("error_categories") or {}).items():
        if not description or str(description).strip().lower() in ("null", "none", ""):
            continue
        db.add(ErrorAnalysisDB(
            student_id=student_record.id if student_record else None,
            course_id=assign.course_id,
            assignment_name=assign.assignment_name,
            error_category=category_label_map.get(key, key.capitalize()),
            error_type=str(description)[:500],
        ))

    db.commit()
    return {
        "status": "success",
        "grade": ai_report.get("score_out_of_100", 0),
        "mode": mode,
        "feedback_tone": feedback_tone,
        "report": ai_report,
        "essay_content": student_text,
        "plagiarism_score": plagiarism_score,
        "plagiarism_matches": plagiarism_matches
    }

async def process_multiple_files(
    assignment_id: int,
    files_data: List[tuple],   # each tuple: (filename: str, content: bytes)
    feedback_tone: str,
    db: Session
) -> List[dict]:
    """
    Process a list of (filename, bytes) and return a list of result dicts.
    Does NOT save submissions – returns results for the caller to handle.
    """
    assign = db.query(AssignmentDB).filter(AssignmentDB.id == assignment_id).first()
    if not assign:
        raise HTTPException(404, "Assignment not found")
    if feedback_tone not in {"formal", "encouraging", "strict"}:
        feedback_tone = "formal"

    results = []
    for filename, content in files_data:
        try:
            # Write to temp file to reuse extract_text
            ext = filename.split('.')[-1].lower()
            temp_path = f"temp_{uuid.uuid4()}.{ext}"
            with open(temp_path, "wb") as f:
                f.write(content)
            try:
                student_text = extract_text(temp_path)
            finally:
                if os.path.exists(temp_path):
                    os.remove(temp_path)

            # Build reference
            if assign.is_model_answer:
                mode = "MODEL"
                reference = f"### MODE: MODEL_ANSWER\nQUESTION:\n{assign.assignment_question}\n\nIDEAL ANSWER:\n{assign.model_answer}"
            else:
                mode = "RUBRIC"
                reference = f"### MODE: RUBRIC\nQUESTION:\n{assign.assignment_question}\n\nRUBRIC:\n{assign.rubric}"

            # AI grading
            ai_report = perform_nlp_grading(
                student_text=student_text,
                mode=mode,
                reference=reference,
                feedback_tone=feedback_tone,
            )

            # Plagiarism check (against existing DB)
            course = db.query(CourseDB).filter(CourseDB.id == assign.course_id).first()
            owner_id = course.user_id if course else 1
            plagiarism_result = await check_plagiarism_ai_direct(
                student_text=student_text,
                user_id=owner_id,
                db=db
            )
            plagiarism_score = plagiarism_result.get("plagiarism_score", 0)
            plagiarism_matches = plagiarism_result.get("matches", [])
            ai_report["plagiarism"] = plagiarism_score
            ai_report["plagiarism_matches"] = plagiarism_matches

            results.append({
                "filename": filename,
                "student_name": filename.split('.')[0],
                "student_text": student_text,
                "ai_grade": ai_report.get("score_out_of_100", 0),
                "plagiarism_score": plagiarism_score,
                "grade_report": ai_report,
                "success": True,
            })
        except Exception as e:
            results.append({
                "filename": filename,
                "success": False,
                "error": str(e),
            })
    return results

@app.post("/grade-submission-batch/{assignment_id}")
async def grade_submission_batch(
    assignment_id: int,
    files: List[UploadFile] = File(...),
    feedback_tone: str = Form("formal"),
    db: Session = Depends(get_db)
):
    # Read all files into memory
    files_data = []
    for f in files:
        content = await f.read()
        files_data.append((f.filename, content))

    results = await process_multiple_files(assignment_id, files_data, feedback_tone, db)

    # Save successful submissions
    assign = db.query(AssignmentDB).filter(AssignmentDB.id == assignment_id).first()
    submissions_to_add = []
    for r in results:
        if r["success"]:
            sub = SubmissionDB(
                assignment_id=assignment_id,
                student_name=r["student_name"],
                status="ready",
                ai_grade=r["ai_grade"],
                plagiarism_score=r["plagiarism_score"],
                grade_report=r["grade_report"],
                essay_content=r["student_text"],
            )
            submissions_to_add.append(sub)
            db.add(sub)

    db.flush()  # get IDs for error analysis

    # Add error analysis entries (same as single grading)
    category_label_map = {
        "conceptual": "Conceptual", "structural": "Structural",
        "language": "Language", "completeness": "Completeness",
    }
    for sub in submissions_to_add:
        student_record = db.query(StudentDB).filter(
            StudentDB.name == sub.student_name,
            StudentDB.course_id == assign.course_id
        ).first()
        for key, description in (sub.grade_report.get("error_categories") or {}).items():
            if not description or str(description).strip().lower() in ("null", "none", ""):
                continue
            db.add(ErrorAnalysisDB(
                student_id=student_record.id if student_record else None,
                course_id=assign.course_id,
                assignment_name=assign.assignment_name,
                error_category=category_label_map.get(key, key.capitalize()),
                error_type=str(description)[:500],
            ))

    db.commit()

    # Return summary
    return {"results": [
        {"filename": r["filename"], "success": r["success"], "grade": r.get("ai_grade"), "error": r.get("error")}
        for r in results
    ]}

@app.post("/analyze-general-submission")
async def analyze_general(data: dict, user_id: int, db: Session = Depends(get_db)):
    student_text = data.get("student_text")
    mode = data.get("mode")
    reference = data.get("reference_content")
    feedback_tone = data.get("feedback_tone", "formal")

    if not all([student_text, mode, reference]):
        raise HTTPException(status_code=400, detail="Missing required fields")

    try:
        report = perform_nlp_grading(
            student_text=student_text,
            mode=mode,
            reference=reference,
            feedback_tone=feedback_tone
        )
        try:
            plagiarism_result = await check_plagiarism_ai_direct(
                student_text=student_text,
                user_id=user_id,
                db=db
            )
            report["plagiarism"] = plagiarism_result.get("plagiarism_score", 0)
            report["plagiarism_matches"] = plagiarism_result.get("matches", [])
            print(f" Plagiarism added to response: {report['plagiarism']}%")
        except Exception as e:
            print(f"Plagiarism check error: {e}")

        return report
    except Exception as e:
        print(f"AI Service Error: {e}")
        raise HTTPException(status_code=500, detail=f"AI grading failed: {str(e)}")

@app.post("/analyze-general-batch")
async def analyze_general_batch(
    user_id: int = Form(...),
    mode: str = Form(...),
    reference_content: str = Form(...),
    feedback_tone: str = Form("formal"),
    files: List[UploadFile] = File(...),
    db: Session = Depends(get_db)
):
    if feedback_tone not in {"formal", "encouraging", "strict"}:
        feedback_tone = "formal"

    results = []
    for file in files:
        try:
            # 1. Extract text from the uploaded file
            content = await file.read()
            ext = file.filename.split('.')[-1].lower()
            temp_path = f"temp_{uuid.uuid4()}.{ext}"
            with open(temp_path, "wb") as f:
                f.write(content)
            try:
                student_text = extract_text(temp_path)
            finally:
                if os.path.exists(temp_path):
                    os.remove(temp_path)

            # 2. Run AI grading (same as single)
            ai_report = perform_nlp_grading(
                student_text=student_text,
                mode=mode,
                reference=reference_content,
                feedback_tone=feedback_tone,
            )

            # 3. Plagiarism check
            plagiarism_result = await check_plagiarism_ai_direct(
                student_text=student_text,
                user_id=user_id,
                db=db
            )
            plagiarism_score = plagiarism_result.get("plagiarism_score", 0)
            plagiarism_matches = plagiarism_result.get("matches", [])
            ai_report["plagiarism"] = plagiarism_score
            ai_report["plagiarism_matches"] = plagiarism_matches

            # 4. Save submission (no assignment_id → general)
            new_sub = SubmissionDB(
                student_name=file.filename.split('.')[0],
                essay_content=student_text,
                ai_grade=ai_report.get("score_out_of_100", 0),
                plagiarism_score=plagiarism_score,
                grade_report=ai_report,
                status="ready",
                user_id=user_id,
                assignment_id=None
            )
            db.add(new_sub)
            db.flush()

            results.append({
                "filename": file.filename,
                "success": True,
                "id": new_sub.id,
                "ai_grade": new_sub.ai_grade,
                "plagiarism_score": plagiarism_score,
                "report": ai_report,
                "student_name": new_sub.student_name,
            })
        except Exception as e:
            results.append({
                "filename": file.filename,
                "success": False,
                "error": str(e),
            })
    db.commit()
    return {"results": results}

@app.post("/submissions/general")
async def create_general_submission(data: dict, db: Session = Depends(get_db)):
    grade_report = data.get("grade_report", {})

    new_sub = SubmissionDB(
        student_name=data.get("student_name", "Unknown"),
        essay_content=data.get("essay_content"),
        ai_grade=data.get("ai_grade"),
        plagiarism_score=data.get("plagiarism_score", 0),
        grade_report=grade_report,
        status=data.get("status", "ready"),
        assignment_id=None,
        user_id=data.get("user_id")
    )
    db.add(new_sub)
    db.commit()
    db.refresh(new_sub)
    return {
        "id": new_sub.id,
        "student_name": new_sub.student_name,
        "ai_grade": new_sub.ai_grade,
        "plagiarism_score": new_sub.plagiarism_score,
        "status": new_sub.status,
        "essay_content": new_sub.essay_content,
        "grade_report": new_sub.grade_report,
        "submission_time": new_sub.submission_time.isoformat() if new_sub.submission_time else None
    }


@app.get("/submissions/general")
async def get_general_submissions(user_id: int, db: Session = Depends(get_db)):
    return db.query(SubmissionDB).filter(
        SubmissionDB.assignment_id == None,
        SubmissionDB.user_id == user_id
    ).order_by(SubmissionDB.id.desc()).all()


# ── Grade Adjustment & Management Core ──────────────────────────────────

@app.put("/api/submissions/{submission_id}")
def update_submission_grade_api(submission_id: int, data: GradeUpdate, db: Session = Depends(get_db)):
    submission = db.query(SubmissionDB).filter(SubmissionDB.id == submission_id).first()
    if not submission:
        raise HTTPException(status_code=404, detail="Submission not found")

    submission.ai_grade = data.ai_grade
    submission.status = data.status
    db.commit()
    return {"message": "Grade updated successfully"}


@app.put("/update-submission-grade/{submission_id}")
async def update_grade(submission_id: int, data: dict, db: Session = Depends(get_db)):
    submission = db.query(SubmissionDB).filter(SubmissionDB.id == submission_id).first()
    if not submission:
        raise HTTPException(status_code=404, detail="Submission not found")

    submission.manual_grade = int(data.get('final_grade', data.get('ai_grade', 0)))
    submission.status = data.get('status', 'graded')

    db.commit()
    db.refresh(submission)
    return {
        "id": submission.id,
        "student_name": submission.student_name,
        "ai_grade": submission.ai_grade,
        "plagiarism_score": submission.plagiarism_score,
        "status": submission.status,
        "essay_content": submission.essay_content,
        "grade_report": submission.grade_report,
        "submitted_at": submission.submission_time.isoformat() if submission.submission_time else None
    }


@app.post("/submissions/{submission_id}/finalize")
async def finalize_submission(submission_id: int, data: FinalizeRequest, db: Session = Depends(get_db)):
    submission = db.query(SubmissionDB).filter(SubmissionDB.id == submission_id).first()
    if not submission:
        raise HTTPException(status_code=404, detail="Submission not found")

    submission.ai_grade = data.manual_grade
    submission.status = "graded"
    db.commit()
    db.refresh(submission)
    return {
        "status": "success",
        "message": "Grade finalized",
        "final_grade": submission.ai_grade
    }


@app.get("/submission-details/{submission_id}")
async def get_submission_details(submission_id: int, db: Session = Depends(get_db)):
    submission = db.query(SubmissionDB).filter(SubmissionDB.id == submission_id).first()
    if not submission:
        raise HTTPException(status_code=404, detail="Submission not found")

    return {
        "id": submission.id,
        "student_name": submission.student_name,
        "ai_grade": submission.ai_grade,
        "plagiarism_score": submission.plagiarism_score,
        "status": submission.status,
        "essay_content": submission.essay_content,
        "submitted_at": submission.submission_time.isoformat() if submission.submission_time else None,
        "grade_report": submission.grade_report or {
            "summary": "Analysis complete.",
            "detected_language": "English"
        }
    }

@app.get("/export-grades-excel/{assignment_id}")
async def export_grades_excel(assignment_id: int, db: Session = Depends(get_db)):
    assignment = db.query(AssignmentDB).filter(AssignmentDB.id == assignment_id).first()
    if not assignment:
        raise HTTPException(404, "Assignment not found")
    submissions = db.query(SubmissionDB).filter(SubmissionDB.assignment_id == assignment_id).all()
    data = []
    for sub in submissions:
        data.append({
            "Student Name": sub.student_name,
            "AI Grade": sub.ai_grade,
            "Manual Grade": sub.manual_grade,
            "Final Grade": sub.manual_grade if sub.manual_grade is not None else sub.ai_grade,
            "Plagiarism Score": sub.plagiarism_score,
            "Submission Date": sub.submission_time.strftime("%Y-%m-%d %H:%M") if sub.submission_time else "",
            "Status": sub.status,
        })
    df = pd.DataFrame(data)
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name="Grades", index=False)
    output.seek(0)
    filename = f"grades_{assignment.assignment_name}.xlsx"
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )
# ──  exporting assignment grades  ────────────────────────────────────────────────────────
@app.get("/export-grades-pdf/{assignment_id}")
async def export_grades_pdf(assignment_id: int, db: Session = Depends(get_db)):
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    assignment = db.query(AssignmentDB).filter(AssignmentDB.id == assignment_id).first()
    if not assignment:
        raise HTTPException(404, "Assignment not found")
    submissions = db.query(SubmissionDB).filter(SubmissionDB.assignment_id == assignment_id).all()
    data = [["Student Name", "AI Grade", "Manual Grade", "Final Grade", "Plagiarism", "Submission Date", "Status"]]
    for sub in submissions:
        final_grade = sub.manual_grade if sub.manual_grade is not None else sub.ai_grade
        data.append([
            sub.student_name or "",
            str(sub.ai_grade) if sub.ai_grade is not None else "",
            str(sub.manual_grade) if sub.manual_grade is not None else "",
            str(final_grade) if final_grade is not None else "",
            str(sub.plagiarism_score) if sub.plagiarism_score is not None else "",
            sub.submission_time.strftime("%Y-%m-%d %H:%M") if sub.submission_time else "",
            sub.status or "",
        ])
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=30)
    styles = getSampleStyleSheet()
    title_style = styles['Title']
    title = Paragraph(f"Grades for Assignment: {assignment.assignment_name}", title_style)

    table = Table(data, repeatRows=1)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.grey),
        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,0), 10),
        ('BOTTOMPADDING', (0,0), (-1,0), 8),
        ('BACKGROUND', (0,1), (-1,-1), colors.beige),
        ('GRID', (0,0), (-1,-1), 1, colors.black),
    ]))
    elements = [title, Spacer(1, 0.2*inch), table]
    doc.build(elements)
    buffer.seek(0)
    filename = f"grades_{assignment.assignment_name}.pdf"
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )
# ── exporting grades of general grading ────────────────────────────────────────────────────────
@app.get("/export-general-excel/{user_id}")
async def export_general_excel(user_id: int, db: Session = Depends(get_db)):
    submissions = db.query(SubmissionDB).filter(
        SubmissionDB.user_id == user_id,
        SubmissionDB.assignment_id == None
    ).all()

    data = []
    for sub in submissions:
        final_grade = sub.manual_grade if sub.manual_grade is not None else sub.ai_grade
        data.append({
            "Student Name": sub.student_name or "",
            "AI Grade": sub.ai_grade,
            "Manual Grade": sub.manual_grade,
            "Final Grade": final_grade,
            "Plagiarism Score": sub.plagiarism_score,
            "Submission Date": sub.submission_time.strftime("%Y-%m-%d %H:%M") if sub.submission_time else "",
            "Status": sub.status or "",
        })
    df = pd.DataFrame(data)
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name="Grades", index=False)
    output.seek(0)
    filename = f"general_grades_user_{user_id}.xlsx"
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@app.get("/export-general-pdf/{user_id}")
async def export_general_pdf(user_id: int, db: Session = Depends(get_db)):
    submissions = db.query(SubmissionDB).filter(
        SubmissionDB.user_id == user_id,
        SubmissionDB.assignment_id == None
    ).all()

    table_data = [["Student Name", "AI Grade", "Manual Grade", "Final Grade", "Plagiarism", "Submission Date", "Status"]]
    for sub in submissions:
        final_grade = sub.manual_grade if sub.manual_grade is not None else sub.ai_grade
        table_data.append([
            sub.student_name or "",
            str(sub.ai_grade) if sub.ai_grade is not None else "",
            str(sub.manual_grade) if sub.manual_grade is not None else "",
            str(final_grade) if final_grade is not None else "",
            str(sub.plagiarism_score) if sub.plagiarism_score is not None else "",
            sub.submission_time.strftime("%Y-%m-%d %H:%M") if sub.submission_time else "",
            sub.status or "",
        ])

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=30)
    styles = getSampleStyleSheet()
    title = Paragraph(f"General Grades (User {user_id})", styles['Title'])

    table = Table(table_data, repeatRows=1)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.grey),
        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,0), 10),
        ('BOTTOMPADDING', (0,0), (-1,0), 8),
        ('BACKGROUND', (0,1), (-1,-1), colors.beige),
        ('GRID', (0,0), (-1,-1), 1, colors.black),
    ]))

    elements = [title, Spacer(1, 0.2*inch), table]
    doc.build(elements)
    buffer.seek(0)

    filename = f"general_grades_user_{user_id}.pdf"
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

# ── Performance Upload ────────────────────────────────────────────────────────

@app.post("/upload-performance")
async def upload_performance_sheet(file: UploadFile = File(...), db: Session = Depends(get_db)):
    content = await file.read()
    try:
        df = pd.read_csv(io.BytesIO(content)) if file.filename.endswith(".csv") \
            else pd.read_excel(io.BytesIO(content), engine='openpyxl')
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid or corrupted file.")

    required = {"student_id", "course_id", "grade", "attendance"}
    if not required.issubset(df.columns):
        raise HTTPException(status_code=400, detail=f"File must contain: {required}")

    added = 0
    for _, row in df.iterrows():
        student = db.query(StudentDB).filter(
            StudentDB.student_id == str(row["student_id"]),
            StudentDB.course_id == int(row["course_id"])
        ).first()
        if student:
            db.add(PerformanceDB(
                student_id=student.id,
                course_id=int(row["course_id"]),
                grade=float(row["grade"]),
                attendance=float(row["attendance"])
            ))
            added += 1
    db.commit()
    return {"message": f"Successfully processed {added} records"}


# ── File Extraction ───────────────────────────────────────────────────────────

@app.post("/extract-text")
async def api_extract_text(file: UploadFile = File(...)):
    temp_path = f"temp_{file.filename}"
    with open(temp_path, "wb") as f:
        shutil.copyfileobj(file.file, f)
    try:
        return {"extracted_text": extract_text(temp_path)}
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


# ── Routers ───────────────────────────────────────────────────────────────────

app.include_router(analysis.router)
app.include_router(lecture.router)
app.include_router(courses_router.router)
